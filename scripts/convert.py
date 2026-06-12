#!/usr/bin/env python3
"""
Convert the Shega/AKOFADA DFS glossary copy decks (.docx) into structured,
machine-readable datasets (JSONL + CSV).

THE ONE RULE: text comes from the editable .docx copy decks, never the PDFs.
Ge'ez (Amharic) corrupts on PDF copy; the PDFs are only for human fidelity
spot-checks. This script reads the .docx sources directly.

FIDELITY POLICY (verified documents — do not alter content):
The decks are expert-verified. The ONLY transformations applied to any stored
text are:
  1. whitespace handling — trim leading/trailing, collapse internal space/tab
     runs, and convert the non-breaking space (U+00A0) to a normal space;
  2. case-insensitive matching — used ONLY for the cross-language join key,
     never written back into the data.
No Unicode (NFC) normalisation, no spelling/word changes, no category
re-casing or merging. Categories are stored exactly as the document wrote them.

Two source formats, one target model:
  - Amharic deck: a single 3-column table (Term | Definition | Example) whose
    rows strictly alternate English -> Amharic. 87 pairs.
  - Afaan Oromo deck: heading/paragraph blocks. Each entry is an English term
    (Heading 1), an Oromo term, a category, then "Hiika:" (definition) and
    "Fakkeenya:" (example). 86 entries.

The decks are complementary and join on term_en (case-insensitive):
  - Amharic deck has English definition + example prose (Oromo deck lacks it).
  - Oromo deck has a category per term (Amharic deck lacks it).
We cross-fill both directions via the English term.

Target record (locked with the team):
  id, term_en, term_local, language, category,
  definition_local, definition_en, example_local, example_en, version

Usage:  python scripts/convert.py
Outputs (under data/): am.jsonl, am.csv, om.jsonl, om.csv
"""

import csv
import io
import json
import re
import sys
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
DATA = ROOT / "data"

AM_DOCX = DOCS / "DFS Translations 4th FINAL.docx"
OM_DOCX = DOCS / "Shega - DFS Glossary in Afaan Oromo - Final.docx"

VERSION = "1.0"

EXPECTED = {"am": 87, "om": 86}

NBSP = " "

# A category in the Oromo deck is the short line between the term and the
# "Hiika"/"Fakkeenya" body. Detected positionally by length, then stored
# VERBATIM (whitespace-trim only) — no canonicalisation, merging, or re-casing.
CATEGORY_MAXLEN = 40


def clean(text):
    """Whitespace-only tidy. Does not alter any characters, words, or case."""
    if text is None:
        return None
    text = text.replace(NBSP, " ")             # non-breaking space -> normal space
    text = re.sub(r"[ \t]+", " ", text)        # collapse internal space/tab runs
    return text.strip()


def is_geez(s):
    return any("ሀ" <= ch <= "፿" for ch in s)


def norm_key(term_en):
    """Cross-language join key: case-insensitive + whitespace only.

    This is the ONLY sanctioned matching transform. It does not touch spelling,
    punctuation, parentheticals, or the stored term text.
    """
    k = (term_en or "").replace(NBSP, " ")
    k = re.sub(r"\s+", " ", k).strip().lower()
    k = re.sub(r"\s*/\s*", "/", k)   # whitespace around a slash is still whitespace
    return k


# --------------------------------------------------------------------------- #
# Amharic: deterministic table parser
# --------------------------------------------------------------------------- #
def parse_amharic(path):
    d = docx.Document(str(path))
    if not d.tables:
        sys.exit(f"[am] expected a table in {path.name}, found none")
    rows = [[c.text for c in r.cells] for r in d.tables[0].rows]
    # Drop the header row ("Term | Definition | Example").
    data = rows[1:]
    if len(data) % 2 != 0:
        print(f"[am] WARNING: {len(data)} data rows is odd; expected EN/AM pairs")

    records = []
    for i in range(0, len(data) - 1, 2):
        en, am = data[i], data[i + 1]
        if is_geez(en[0]) or not is_geez(am[0]):
            print(f"[am] WARNING: row pair {i} breaks EN->AM alternation: "
                  f"{en[0][:30]!r} / {am[0][:30]!r}")
        records.append({
            "term_en": clean(en[0]),
            "term_local": clean(am[0]),
            "language": "am",
            "category": None,                  # not in this deck; backfilled later
            "definition_local": clean(am[1]),
            "definition_en": clean(en[1]),
            "example_local": clean(am[2]) or None,
            "example_en": clean(en[2]) or None,
            "version": VERSION,
        })
    return records


# --------------------------------------------------------------------------- #
# Afaan Oromo: heading/label parser
# --------------------------------------------------------------------------- #
def _logical_lines(paras, start, end):
    """Flatten an entry's paragraphs into logical lines (split embedded \\n)."""
    out = []
    for text, style in paras[start:end]:
        for piece in text.split("\n"):
            piece = piece.strip()
            if piece:
                out.append((piece, style))
    return out


def _strip_label(line, label):
    """If line starts with `label` (optionally + ':'), return the remainder."""
    m = re.match(rf"^{label}\s*:?\s*(.*)$", line, flags=re.IGNORECASE | re.DOTALL)
    return m.group(1).strip() if m else None


def _looks_like_category(line):
    """A category is a short line with no sentence-ending punctuation, sitting
    between the term and the Hiika/Fakkeenya body. Detected by shape only; the
    value itself is preserved verbatim by the caller."""
    if len(line) > CATEGORY_MAXLEN:
        return False
    # definitions/examples are full sentences; categories are short noun labels
    if line.endswith((".", "::", "።")):
        return False
    return True


def parse_oromo(path):
    d = docx.Document(str(path))
    paras = [(p.text.strip(), (p.style.name if p.style else ""))
             for p in d.paragraphs if p.text.strip()]
    n = len(paras)

    # Heading runs delimit entries: first heading = English term; a 2nd
    # consecutive heading (if present) = the Oromo term.
    runs = []
    i = 0
    while i < n:
        if paras[i][1] == "Heading 1":
            j = i
            while j < n and paras[j][1] == "Heading 1":
                j += 1
            runs.append((i, j))   # [start, end) of the heading run
            i = j
        else:
            i += 1

    records = []
    for idx, (h_start, h_end) in enumerate(runs):
        term_en = paras[h_start][0]
        run_len = h_end - h_start
        # Entry content spans up to the next entry's first heading.
        block_end = runs[idx + 1][0] if idx + 1 < len(runs) else n

        if run_len >= 2:
            term_local = paras[h_start + 1][0]
            lines = _logical_lines(paras, h_start + 2, block_end)
        else:
            # Oromo term is the first logical line after the English heading.
            lines = _logical_lines(paras, h_start + 1, block_end)
            term_local = lines[0][0] if lines else ""
            lines = lines[1:]

        category = None
        definition = []
        example = []
        bucket = None  # None | 'def' | 'ex'
        for text, _style in lines:
            rest = _strip_label(text, "Fakkeenya")
            if rest is not None:
                bucket = "ex"
                if rest:
                    example.append(rest)
                continue
            rest = _strip_label(text, "Hiika")
            if rest is not None:
                bucket = "def"
                if rest:
                    definition.append(rest)
                continue
            rest = _strip_label(text, "Ibsa")
            if rest is not None:
                bucket = "def"          # description folds into definition
                if rest:
                    definition.append(rest)
                continue
            if bucket is None and category is None and _looks_like_category(text):
                category = text          # stored VERBATIM (clean() applied later)
                continue
            # Unlabeled body text. Several entries give the definition as a
            # plain paragraph with no "Hiika:" label, so default to the
            # definition bucket until a "Fakkeenya:" label switches us to example.
            if bucket is None:
                bucket = "def"
            if bucket == "ex":
                example.append(text)
            else:
                definition.append(text)

        records.append({
            "term_en": clean(term_en),
            "term_local": clean(term_local),
            "language": "om",
            "category": clean(category),
            "definition_local": clean(" ".join(definition)) or None,
            "definition_en": None,             # not in this deck; backfilled later
            "example_local": clean(" ".join(example)) or None,
            "example_en": None,                # not in this deck; backfilled later
            "version": VERSION,
        })
    return records


# --------------------------------------------------------------------------- #
# Cross-fill on term_en (case-insensitive), assign IDs, write, validate
# --------------------------------------------------------------------------- #
def cross_fill(am, om):
    en_prose = {norm_key(r["term_en"]): (r["definition_en"], r["example_en"])
                for r in am}
    cat_by_en = {norm_key(r["term_en"]): r["category"] for r in om if r["category"]}

    am_filled = 0
    for r in am:
        cat = cat_by_en.get(norm_key(r["term_en"]))
        if cat and not r["category"]:
            r["category"] = cat
            am_filled += 1

    om_filled = 0
    for r in om:
        prose = en_prose.get(norm_key(r["term_en"]))
        if prose:
            d_en, e_en = prose
            if d_en and not r["definition_en"]:
                r["definition_en"] = d_en
            if e_en and not r["example_en"]:
                r["example_en"] = e_en
            om_filled += 1

    print(f"[join] backfilled category onto {am_filled}/{len(am)} Amharic entries")
    print(f"[join] backfilled English prose onto {om_filled}/{len(om)} Oromo entries")

    am_keys = {norm_key(r["term_en"]) for r in am}
    om_keys = {norm_key(r["term_en"]) for r in om}
    only_om = sorted(k for k in om_keys - am_keys)
    only_am = sorted(k for k in am_keys - om_keys)
    if only_am:
        print(f"[join] {len(only_am)} terms only in Amharic (no Oromo match, "
              f"not cross-filled): {only_am}")
    if only_om:
        print(f"[join] {len(only_om)} terms only in Oromo (no Amharic match, "
              f"not cross-filled): {only_om}")


def assign_ids(records, lang):
    for i, r in enumerate(records, start=1):
        r["id"] = f"{lang}-{i:03d}"
    order = ["id", "term_en", "term_local", "language", "category",
             "definition_local", "definition_en", "example_local",
             "example_en", "version"]
    return [{k: r[k] for k in order} for r in records]


def write_outputs(records, lang):
    DATA.mkdir(exist_ok=True)
    jsonl = DATA / f"{lang}.jsonl"
    with io.open(jsonl, "w", encoding="utf-8") as f:
        for r in records:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")
    csvp = DATA / f"{lang}.csv"
    with io.open(csvp, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(records[0].keys()))
        w.writeheader()
        for r in records:
            w.writerow(r)
    print(f"[write] {jsonl.name} + {csvp.name} ({len(records)} records)")


def validate(records, lang):
    problems = []
    if len(records) != EXPECTED[lang]:
        problems.append(f"count {len(records)} != expected {EXPECTED[lang]}")
    ids = [r["id"] for r in records]
    if len(set(ids)) != len(ids):
        problems.append("duplicate ids")
    required = ["id", "term_en", "term_local", "language",
                "definition_local", "version"]
    for r in records:
        for field in required:
            if not r.get(field):
                problems.append(f"{r['id']}: empty required field {field!r}")
    for r in records:
        json.loads(json.dumps(r, ensure_ascii=False))  # round-trips as UTF-8 JSON
    cats = sum(1 for r in records if r["category"])
    ex_local = sum(1 for r in records if r["example_local"])
    def_en = sum(1 for r in records if r["definition_en"])
    print(f"[validate:{lang}] {len(records)} records | category {cats} | "
          f"example_local {ex_local} | definition_en {def_en}")
    if problems:
        print(f"[validate:{lang}] ISSUES:")
        for p in problems:
            print("   -", p)
    else:
        print(f"[validate:{lang}] OK")
    return not problems


def main():
    if not AM_DOCX.exists():
        sys.exit(f"missing {AM_DOCX}")
    if not OM_DOCX.exists():
        sys.exit(f"missing {OM_DOCX}")

    am = parse_amharic(AM_DOCX)
    om = parse_oromo(OM_DOCX)
    print(f"[parse] amharic={len(am)}  oromo={len(om)}")

    cross_fill(am, om)

    am = assign_ids(am, "am")
    om = assign_ids(om, "om")

    write_outputs(am, "am")
    write_outputs(om, "om")

    ok_am = validate(am, "am")
    ok_om = validate(om, "om")

    total = len(am) + len(om)
    print(f"[done] {total} total entries (target 173)")
    sys.exit(0 if (ok_am and ok_om) else 1)


if __name__ == "__main__":
    main()
